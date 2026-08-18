#!/usr/bin/env python3
"""테스터 확인·질문지 (워드). 문항은 tester_survey_spec.py 한 곳에서 온다.

쓰는 법:  python3 tool/tester_survey/make_tester_doc.py
자세한 것은 같은 폴더의 README.md.
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import tester_survey_spec as S
from assets_dir import assets_path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = assets_path('문서', '커넥션센스_테스터_확인및질문지_2026-08-18.docx')
FONT = 'Apple SD Gothic Neo'

d = Document()
for sec in d.sections:
    sec.top_margin = sec.bottom_margin = Cm(1.8)
    sec.left_margin = sec.right_margin = Cm(2.0)
st = d.styles['Normal']
st.font.name = FONT
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(5)
st.paragraph_format.line_spacing = 1.3


def H(t, level=1):
    p = d.add_heading(t, level=level)
    for r in p.runs:
        r.font.name = FONT
        r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    return p


def P(t='', bold=False, size=10.5, indent=0.0, after=5):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    for i, line in enumerate(t.split('\n')):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = FONT
    return p


def TABLE(headers, rows, widths=None):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.name = FONT
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
            r.font.name = FONT
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    d.add_paragraph().paragraph_format.space_after = Pt(3)
    return t


def LINES(n=3, label=''):
    if label:
        P(label, bold=True, after=2)
    for _ in range(n):
        p = P('_' * 78, size=9, after=8)
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


CHK = '☐ 잘 됨   ☐ 안 됨   ☐ 안 해봄'
RATE = '◎  ○  △  ✗  −'

# ══════════════ 표지 + 기기 환경
t = d.add_heading(S.TITLE, level=0)
for r in t.runs:
    r.font.name = FONT
p = P(S.SUBTITLE, size=10)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

P(S.INTRO, size=10)
P(S.INTRO_WARN, bold=True, size=10)

H('테스트 환경', 1)
TABLE(['항목', '적어 주세요'],
      [[label, '' if not hint else f'({hint})'] for label, _, hint in S.HEAD],
      widths=[4.5, 11.5])

H(S.PREP_TITLE, 1)
P(S.PREP)

d.add_page_break()

# ══════════════ 1부 — 화면 순서
H('1부. 화면별 확인 — 앱을 여는 순서대로', 1)
P('적힌 대로 안 되면 그것이 바로 제보 대상입니다. ⭐ 표시는 이번에 새로 넣거나 고친 것입니다.')

for title, note, rows, free in S.SCREENS:
    H(title, 2)
    if note:
        P(note, bold=note.startswith('⚠️'))
    if title.startswith('4.'):
        for q, choices, help_ in S.AI_Q:
            P(q, bold=True)
            if help_:
                P(help_, indent=0.6, size=10)
            P('   '.join('☐ ' + c for c in choices), indent=0.6)
            P()
        LINES(2, S.AI_FREE[0])
        P(S.AI_FREE[1], size=10)
        continue
    if rows:
        TABLE(['확인해 주실 것', '결과', '메모'],
              [[r + ('\n' + h if h else ''), CHK, ''] for r, h in rows],
              widths=[7.4, 4.4, 4.2])
    if free and free[0]:
        LINES(3 if '자유롭게' in free[0] else 2, free[0])
        if len(free) > 1 and free[1]:
            P(free[1], bold=True, size=10)
    if title.startswith('1.'):
        P()
        P(S.ACC_TITLE, bold=True)
        TABLE(['칸', '정확도', '칸', '정확도'], [list(r) for r in S.ACC_ROWS],
              widths=[4.0, 3.0, 4.6, 3.0])
        P('⚠️ 회사와 직함은 아직 낮습니다. 고쳐 쓰셔야 할 때가 많을 것입니다.', bold=True)

d.add_page_break()

# ══════════════ 2부 — 만족도
H('2부. ' + S.SURVEY_TITLE, 1)
P(S.SURVEY_NOTE)
P('표시: ◎ 해결됨 · ○ 나아졌지만 아쉬움 · △ 그대로 · ✗ 더 나빠짐 · − 기억 안 남 / 안 써봄',
  bold=True, size=10)

H('오류 (E)', 2)
TABLE(['#', '알려주신 내용', '저희가 한 것', '평가', '한 줄'],
      [[a, b, c, RATE, ''] for a, b, c in S.E_ITEMS],
      widths=[1.4, 4.6, 5.4, 2.6, 3.0])

d.add_page_break()
H('기능 개선 (F)', 2)
P('표시: ◎ 해결됨 · ○ 나아졌지만 아쉬움 · △ 그대로 · ✗ 더 나빠짐 · − 기억 안 남 / 안 써봄',
  bold=True, size=10)
TABLE(['#', '알려주신 내용', '저희가 한 것', '평가', '한 줄'],
      [[a, b, c, RATE, ''] for a, b, c in S.F_ITEMS],
      widths=[1.4, 4.6, 5.4, 2.6, 3.0])

LINES(3, '△(그대로)나 ✗(더 나빠짐)를 고르신 항목이 있으면, 어느 것인지와 어떤 점이 그런지')

P(S.OVERALL_Q[0], bold=True)
P('   '.join('☐ ' + c for c in S.OVERALL_Q[1]), indent=0.6)
LINES(2, S.OVERALL_FREE)

d.add_page_break()

# ══════════════ 3부 — 값
H('3부. ' + S.PRICE_TITLE, 1)
P(S.PRICE_NOTE)
P()
for q, kind, help_, choices in S.PRICE:
    P(q, bold=True)
    if help_:
        P(help_, indent=0.6, size=9.5)
    if kind == 'choice':
        P('     '.join('☐ ' + c for c in choices), indent=0.6)
        P()
    elif kind == 'text':
        LINES(1)
    else:
        LINES(2)

d.add_page_break()

# ══════════════ 4부 — 새로 발견한 것
H('4부. ' + S.NEW_TITLE, 1)
P(S.NEW_NOTE)

H('4-1. 오류 · 이상한 동작', 2)
TABLE(['#', '어느 화면 · 무엇을 했을 때', '무엇이 잘못됐나', '얼마나 자주'],
      [[str(i), '', '', '☐ 항상  ☐ 가끔  ☐ 한 번'] for i in range(1, 7)],
      widths=[1.0, 6.0, 5.4, 4.0])

H('4-2. 있으면 쓰겠다 싶은 기능', 2)
TABLE(['#', '어떤 기능', '어떤 상황에서 필요한가'],
      [[str(i), '', ''] for i in range(1, 6)], widths=[1.0, 6.5, 8.9])

H('4-3. 없어도 될 것 같은 기능', 2)
TABLE(['#', '어떤 기능', '왜 그렇게 느끼셨나'],
      [[str(i), '', ''] for i in range(1, 4)], widths=[1.0, 6.5, 8.9])

H('4-4. 마지막으로', 2)
P(S.LAST_Q[0], bold=True)
P('     '.join('☐ ' + c for c in S.LAST_Q[1]), indent=0.6)
LINES(3, '그렇게 느끼신 이유를 자유롭게 적어 주세요')

d.add_paragraph()
P('고맙습니다. 적어 주신 내용은 그대로 다음 개선에 반영됩니다.', bold=True, size=10)
f = P('보내실 곳: ____________________________', size=9.5)
f.alignment = WD_ALIGN_PARAGRAPH.RIGHT

d.save(OUT)
print('만들었습니다:', OUT)
print('  문단', len(d.paragraphs), '· 표', len(d.tables))
